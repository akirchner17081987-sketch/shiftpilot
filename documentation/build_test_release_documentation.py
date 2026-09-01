from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "documentation" / "SchichtFunk_Test-und-Freigabedokumentation_2026-09-01.docx"
LOGO = ROOT / "assets" / "schichtfunk-company-logo.png"

NAVY = "0B1D2A"
TEAL = "1DD6C0"
BLUE = "2E74B5"
INK = "172B3A"
MUTED = "5D6D78"
LIGHT = "F2F4F7"
PALE_TEAL = "E8FAF7"
PALE_GOLD = "FFF5D6"
PALE_RED = "FDECEC"
GREEN = "18794E"
GOLD = "8A6200"
RED = "A12626"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_in):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, size=11, bold=False, color=INK, after=6, before=0, italic=False, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    set_run(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_bullet(doc, label, detail, color=INK):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.10
    set_run(p.add_run(label + ": "), size=10.5, bold=True, color=color)
    set_run(p.add_run(detail), size=10.5, color=color)
    return p


def add_status_callout(doc, title, body, fill, accent):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_width(cell, 6.5)
    set_cell_margins(cell, 150, 180, 150, 180)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(title), size=11.5, bold=True, color=accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.08
    set_run(p2.add_run(body), size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(text), size={1:16, 2:13, 3:12}[level], bold=True, color=BLUE if level < 3 else "1F4D78")
    return p


def add_test_table(doc, rows):
    widths = [0.48, 2.20, 2.54, 1.28]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    headers = ["Nr.", "Prüfgegenstand", "Ergebnis und Bewertung", "Status"]
    for i, (cell, text) in enumerate(zip(table.rows[0].cells, headers)):
        set_cell_width(cell, widths[i])
        set_cell_margins(cell)
        shade(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(text), size=9, bold=True, color=WHITE)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for num, criterion, result, status in rows:
        cells = table.add_row().cells
        values = [num, criterion, result, status]
        for i, (cell, value) in enumerate(zip(cells, values)):
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            color = GREEN if status in ("Bestanden", "Korrigiert") else GOLD if status in ("Manuell", "Akzeptiert") else RED
            set_run(p.add_run(value), size=8.7, bold=(i in (0,3)), color=(color if i == 3 else INK))
            if status in ("Manuell", "Akzeptiert"):
                shade(cell, PALE_GOLD)
            elif status in ("Bestanden", "Korrigiert") and i == 3:
                shade(cell, PALE_TEAL)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_section_page(doc, number, title, purpose, rows, evidence, conclusion):
    # Start major blocks on a fresh page, but let sections 5 and 6 use the
    # remaining space left by their long predecessor tables.
    if number in {1, 2, 3, 4, 8}:
        doc.add_page_break()
    add_heading(doc, f"{number}. {title}", 1)
    add_text(doc, purpose, size=10.5, color=MUTED, after=9)
    add_test_table(doc, rows)
    add_heading(doc, "Nachweise", 2)
    for label, detail in evidence:
        add_bullet(doc, label, detail)
    add_status_callout(doc, "Teilbewertung", conclusion, PALE_TEAL, GREEN)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.82)
section.bottom_margin = Inches(0.78)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.42)
section.footer_distance = Inches(0.42)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, before, after, color in (("Heading 1",16,16,8,BLUE),("Heading 2",13,12,6,BLUE),("Heading 3",12,8,4,"1F4D78")):
    s = styles[name]
    s.font.name = "Calibri"
    s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)

# Running header/footer
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(hp.add_run("SCHICHTFUNK  |  Test- und Freigabedokumentation"), size=8.5, bold=True, color=MUTED)
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(fp.add_run("Vertrauliche Projektdokumentation  •  Stand 01.09.2026"), size=8, color=MUTED)

# Cover
add_text(doc, "SCHICHTFUNK", size=11, bold=True, color=TEAL, after=8, before=18)
if LOGO.exists():
    p_logo = doc.add_paragraph()
    p_logo.paragraph_format.space_after = Pt(22)
    p_logo.add_run().add_picture(str(LOGO), width=Inches(2.2))
add_text(doc, "TEST- UND\nFREIGABEDOKUMENTATION", size=25, bold=True, color=NAVY, after=10)
add_text(doc, "Mitarbeiterportal – vollständige Prüfung nach Schema Punkt 1 bis Punkt 8", size=14, color=BLUE, after=28)

meta = doc.add_table(rows=5, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta.autofit = False
for idx, (label, value) in enumerate([
    ("Dokumenttyp", "Abnahme- und Produktionsfreigabebericht"),
    ("Prüfobjekt", "SchichtFunk Mitarbeiterportal"),
    ("Produktionsadresse", "https://shiftpilot-two.vercel.app/"),
    ("Prüfstand", "01.09.2026"),
    ("Gesamtstatus", "Produktiv freigegeben mit dokumentierten Restrisiken"),
]):
    for j, width in enumerate((1.7, 4.8)):
        set_cell_width(meta.rows[idx].cells[j], width)
        set_cell_margins(meta.rows[idx].cells[j], 100, 140, 100, 140)
    shade(meta.rows[idx].cells[0], LIGHT)
    set_run(meta.rows[idx].cells[0].paragraphs[0].add_run(label), size=9.5, bold=True, color=NAVY)
    color = GREEN if label == "Gesamtstatus" else INK
    set_run(meta.rows[idx].cells[1].paragraphs[0].add_run(value), size=9.5, bold=(label == "Gesamtstatus"), color=color)

add_text(doc, "Dokumentzweck", size=12, bold=True, color=BLUE, before=24, after=5)
add_text(doc, "Diese Dokumentation fasst die fachlichen, datenschutzbezogenen, technischen und betrieblichen Prüfungen des Mitarbeiterportals zusammen. Sie trennt automatisierte Nachweise, manuelle Abnahmen, behobene Fehler und ausdrücklich akzeptierte Restrisiken.", size=10.5, color=MUTED, after=10)
add_status_callout(doc, "Freigabeaussage", "Die produktive Veröffentlichung wurde abgeschlossen. Die Freigabe gilt unter Kenntnis der drei in Abschnitt 8 dokumentierten und ausdrücklich akzeptierten Restrisiken.", PALE_TEAL, GREEN)

doc.add_page_break()
add_heading(doc, "Management-Zusammenfassung", 1)
add_status_callout(doc, "Gesamturteil: FREIGEGEBEN MIT RESTRISIKEN", "Die priorisierten Fehler wurden korrigiert, die zentralen Portalabläufe mit Mitarbeiter- und Manager-Testkonto abgenommen und die Produktion technisch geprüft. Acht öffentliche End-to-End-Tests liefen erfolgreich; acht geschützte End-to-End-Tests wurden mangels hinterlegter Testkonto-Secrets übersprungen und durch manuelle Abnahme abgedeckt.", PALE_TEAL, GREEN)
add_heading(doc, "Statuslegende", 2)
for label, detail in [
    ("Bestanden", "Prüfung ohne festgestellten Fehler abgeschlossen."),
    ("Korrigiert", "Fehler gefunden, behoben, erneut geprüft und bestanden."),
    ("Manuell", "Ablauf im angemeldeten Portal mit Testkonto geprüft."),
    ("Akzeptiert", "Offener Punkt wurde bewusst als Restrisiko für diese Freigabe akzeptiert."),
]: add_bullet(doc, label, detail)
add_heading(doc, "Prüfumfang", 2)
for item in [
    "11 getrennte Mitarbeiterbereiche sowie die zugehörigen Managerabläufe",
    "Rollen- und Datentrennung über Supabase Row Level Security",
    "Desktop- und Mobilansicht, Tastaturbedienung, Fokusführung und verständliche Rückmeldungen",
    "Build, statische Regressionstests, Browserprüfungen, Abhängigkeiten und Sicherheitsheader",
    "Produktive Bereitstellung auf Vercel einschließlich dokumentiertem Rückfallziel",
]: add_bullet(doc, "Umfang", item)
add_heading(doc, "Priorisierte Prüfreihenfolge", 2)
overview = [
    "Schichttausch-Anzeigefehler beheben", "Datenschutz und Mitarbeitertrennung absichern", "Echte Testzugänge einrichten", "Alle Arbeitsabläufe vollständig testen", "Fehler- und Sonderfälle prüfen", "Mobilgeräte und Barrierefreiheit finalisieren", "Performance und technische Qualität verbessern", "Abnahme und Produktion",
]
for i, item in enumerate(overview, 1): add_bullet(doc, str(i), item)

add_section_page(doc, 1, "Schichttausch-Anzeigefehler beheben",
    "Ziel war eine eindeutige Zuordnung der Schichttausch-Funktion und eine erneute Prüfung aller Portalbereiche auf vermischte Inhalte.",
    [
        ("1.1", "Schichttausch-Karte nur im Bereich „Schichttausch“ anzeigen", "Fehlerhafte Mehrfachdarstellung entfernt; die Karte erscheint ausschließlich im vorgesehenen Bereich.", "Korrigiert"),
        ("1.2", "Alle 11 Portalbereiche auf getrennte Darstellung prüfen", "Bereiche einzeln aufgerufen; Inhalte, Navigation und Zustände bleiben voneinander getrennt.", "Bestanden"),
    ],
    [("Funktionsprüfung", "Schichttausch-Bereich nach Korrektur erneut geprüft."), ("Regression", "Alle 11 Mitarbeiterbereiche nacheinander kontrolliert."), ("Ergebnis", "Keine erneut auftretende Schichttausch-Karte außerhalb des Zielbereichs.")],
    "Punkt 1 ist abgeschlossen. Der ursprüngliche Anzeigefehler wurde behoben und die Trennung der Portalbereiche bestätigt.")

add_section_page(doc, 2, "Datenschutz und Mitarbeitertrennung absichern",
    "Ziel war, dass angemeldete Mitarbeiter ausschließlich ihre eigenen personenbezogenen und arbeitsbezogenen Daten erhalten und verändern können.",
    [
        ("2.1", "Arbeitszeiten, Genehmigungen, Anträge und Abwesenheiten ausschließlich für den angemeldeten Mitarbeiter laden", "Personenbezug über die aktive Sitzung und Mitarbeiterzuordnung geprüft.", "Bestanden"),
        ("2.2", "Unangemessen breite Abfragen aus dem Portal entfernen", "Portalabfragen auf den angemeldeten Benutzer eingeschränkt; keine ungezielte Gesamtabfrage im Mitarbeiterfluss festgestellt.", "Korrigiert"),
        ("2.3", "Supabase-RLS-Regeln vollständig prüfen und versioniert hinterlegen", "Alle 31 öffentlichen Tabellen besitzen aktivierte RLS; anonyme Rollen haben keinen Tabellenzugriff; keine trivial offenen oder veralteten Regeln gefunden.", "Bestanden"),
        ("2.4", "Mit zwei Test-Mitarbeitern den gegenseitigen Datenzugriff ausschließen", "Getrennte Konten und Sitzungen geprüft; wechselseitiger Zugriff auf persönliche Portaldaten war nicht möglich.", "Manuell"),
    ],
    [("RLS-Nachweis", "31/31 öffentliche Tabellen mit aktivierter Row Level Security."), ("Anonymer Zugriff", "Kein Tabellenzugriff für die anonyme Rolle."), ("Kontentrennung", "Zwei Mitarbeiterkonten mit unterschiedlichen Datenbeständen geprüft."), ("Repository", "Datenbank- und Richtlinienänderungen sind im Projektstand versioniert.")],
    "Punkt 2 ist abgeschlossen. Die geprüfte Daten- und Rollentrennung verhindert einen gegenseitigen Zugriff der Mitarbeiterkonten.")

add_section_page(doc, 3, "Echte Testzugänge einrichten",
    "Ziel war eine reproduzierbare Abnahmeumgebung ohne reale Mitarbeiterdaten und mit allen benötigten Rollen.",
    [
        ("3.1", "Separate Testfirma beziehungsweise Testdaten verwenden", "Abnahme erfolgte mit eigens angelegten Testdatensätzen; reale Mitarbeiterdaten wurden nicht verwendet.", "Bestanden"),
        ("3.2", "Mindestens zwei Mitarbeiter und einen Manager anlegen", "Zwei Mitarbeiterkonten sowie ein Manager-/Administratorkonto standen für Rollen- und Ablaufprüfungen bereit.", "Bestanden"),
        ("3.3", "Keine realen Mitarbeiterdaten für Abnahmetests verwenden", "Verwendete Personen, Schichten, Abwesenheiten und Zeiten waren als Testdaten gekennzeichnet.", "Bestanden"),
    ],
    [("Mitarbeiterkonto", "Testkonto „Test Marktplatz 2“ für Portal- und Marktplatzabläufe."), ("Managerkonto", "Testkonto „Alexander Kirchner“ mit Administratorrolle."), ("Zusätzliche Trennung", "Zweites Mitarbeiterkonto zur negativen Zugriffskontrolle."), ("Datenschutz", "Keine Passwörter, Tokens oder persönlichen Zugangsdaten werden in diesem Bericht dokumentiert.")],
    "Punkt 3 ist abgeschlossen. Die benötigten Rollen und getrennten Testdaten standen für die Freigabeprüfung bereit.")

add_section_page(doc, 4, "Alle Arbeitsabläufe vollständig testen",
    "Ziel war die vollständige fachliche Abnahme der wichtigsten Mitarbeiter- und Managerabläufe vom Zugang bis zur Genehmigung.",
    [
        ("4.1", "Anmeldung, Abmeldung und Passwortbehandlung", "An- und Abmeldung geprüft. Fehlerhafte Recovery-Weiterleitung („Auth session missing“) korrigiert. Rate-Limit-Fall erkannt; administrativ angelegtes Testkonto für die Abnahme verwendet.", "Korrigiert"),
        ("4.2", "Schichten anzeigen und bestätigen", "Veröffentlichte Schichten wurden im Mitarbeiterkonto angezeigt und konnten im vorgesehenen Statusfluss bestätigt werden.", "Manuell"),
        ("4.3", "Schichtänderungen und Schichttausch", "Änderungs- und Tauschangebote einschließlich Ablehnung und Statusaktualisierung geprüft.", "Manuell"),
        ("4.4", "Marktplatz und Ersatzanfragen", "Marktplatzangebot und Übernahmeprozess zwischen Mitarbeiter- und Manageransicht konsistent geprüft.", "Manuell"),
        ("4.5", "Arbeitszeiterfassung", "Erfassung und Anzeige von Arbeitszeit einschließlich bestätigtem 9,5-Stunden-Datensatz geprüft.", "Manuell"),
        ("4.6", "Abwesenheitsanträge", "Antrag, Anzeige und genehmigter Zustand für den 10.09.2026 geprüft.", "Manuell"),
        ("4.7", "Stundenkonto und Lohnvorschau", "Berechnung und persönliche Anzeige im Mitarbeiterkonto geprüft; Warnlogik greift erst bei mehr als 10 Stunden.", "Bestanden"),
        ("4.8", "Profil und persönliche Daten", "Profilansicht und persönliche Felder auf Sichtbarkeit, Bearbeitbarkeit und Kontozuordnung geprüft.", "Manuell"),
        ("4.9", "Managerfreigaben, Ablehnungen und Benachrichtigungen", "Genehmigung, Ablehnung und daraus entstehende Mitarbeiterstatusmeldungen kontoübergreifend geprüft.", "Manuell"),
    ],
    [("Gemeinsame Abnahme", "Mitarbeiter- und Managerkonto wurden nacheinander im produktionsnahen Ablauf verwendet."), ("Konsistenz", "Marktplatz, veröffentlichte Schichten, genehmigte Abwesenheit und bestätigte Zeit waren in beiden Rollen konsistent."), ("Datenintegrität", "Während der abschließenden Abnahme wurden keine produktiven Testdaten unbeabsichtigt verändert.")],
    "Punkt 4 ist fachlich abgeschlossen. Die Kernabläufe wurden manuell mit beiden Rollen abgenommen; der Passwort-Recovery-Fehler wurde zuvor korrigiert.")

add_section_page(doc, 5, "Fehler- und Sonderfälle prüfen",
    "Ziel war ein nachvollziehbares Verhalten bei Verbindungs-, Sitzungs-, Daten- und Zeitgrenzen sowie bei konkurrierenden Anträgen.",
    [
        ("5.1", "Keine Netzwerkverbindung", "Offline-/Fehlerzustände liefern eine verständliche Rückmeldung und lassen keine irreführende Erfolgsanzeige stehen.", "Bestanden"),
        ("5.2", "Abgelaufene Anmeldung", "Ungültige oder fehlende Sitzung führt zurück in den Zugang und zeigt eine verständliche Statusmeldung.", "Korrigiert"),
        ("5.3", "Doppelte Anträge", "Doppelte Vorgänge werden nicht als unabhängige erfolgreiche Anträge weitergeführt; Status und Sperren geprüft.", "Bestanden"),
        ("5.4", "Überschneidende Abwesenheiten", "Überlappende Zeiträume werden erkannt beziehungsweise verständlich abgewiesen.", "Bestanden"),
        ("5.5", "Bereits vergebene Tauschangebote", "Nicht mehr verfügbare Angebote können nicht erneut verbindlich übernommen werden; Status wird aktualisiert.", "Bestanden"),
        ("5.6", "Fehlende oder unvollständige Mitarbeiterdaten", "Oberfläche bleibt bedienbar und verwendet verständliche Ersatz- beziehungsweise Fehlhinweise.", "Korrigiert"),
        ("5.7", "Zeitzonen- und Datumswechsel", "Datumsdarstellung und Tageszuordnung an Grenzstellen geprüft; keine fehlerhafte Verschiebung festgestellt.", "Bestanden"),
        ("5.8", "Schichten über Mitternacht", "Beginn, Ende und Dauer werden über den Tageswechsel hinweg korrekt dargestellt und verarbeitet.", "Bestanden"),
    ],
    [("Sitzungsfehler", "Recovery-URL mit fehlender Auth-Sitzung reproduziert und behoben."), ("Statusmeldungen", "Technische Rohmeldungen wurden in verständliche Nutzertexte überführt, soweit im Portal sichtbar."), ("Grenzfälle", "Zeit-, Angebots- und Antragszustände wurden gezielt gegen ungültige Wiederholung geprüft.")],
    "Punkt 5 ist abgeschlossen. Kritische Sonderfälle reagieren kontrolliert; der zuvor sichtbare Sitzungsfehler wurde korrigiert.")

add_section_page(doc, 6, "Mobilgeräte und Barrierefreiheit finalisieren",
    "Ziel war eine mobile, tastaturbedienbare und verständliche Portaloberfläche mit klaren Zuständen und Formularen.",
    [
        ("6.1", "Echte Smartphones sowie Hoch- und Querformat prüfen", "Responsive Layouts auf mobilen Größen und in Hoch-/Querformat geprüft; zentrale Bereiche bleiben lesbar und bedienbar.", "Bestanden"),
        ("6.2", "Tastaturbedienung, Fokusführung, Kontraste und Fehlermeldungen testen", "Interaktive Elemente, sichtbarer Fokus, Lesbarkeit und Fehlerzustände geprüft und an festgestellten Stellen korrigiert.", "Korrigiert"),
        ("6.3", "Formulare mit verständlichen Beschriftungen und Statusmeldungen versehen", "Formularlabels, Hilfetexte, Erfolgs-, Warn- und Fehlermeldungen vollständig geprüft und vereinheitlicht.", "Korrigiert"),
    ],
    [("Mobile Automation", "Öffentliche End-to-End-Prüfungen liefen zusätzlich mit Pixel-7-Geräteprofil."), ("Manuelle Sichtprüfung", "Header, Navigation, Formulare und Karten auf schmalen Ansichten kontrolliert."), ("Barrierefreiheit", "Fokus, Tastaturweg, Kontrast und verständliche Statuskommunikation in die Abnahme einbezogen.")],
    "Punkt 6 ist abgeschlossen. Mobile Darstellung und zentrale Zugänglichkeitsmerkmale wurden geprüft; Beschriftungen und Statusmeldungen wurden final korrigiert.")

add_section_page(doc, 7, "Performance und technische Qualität verbessern",
    "Ziel war eine robuste, wartbare und automatisiert prüfbare Portalversion ohne unnötige Last oder stilles Fehlverhalten.",
    [
        ("7.1", "Nicht benötigte Module erst bei Bedarf laden", "Ladeverhalten und Seitenaufbau geprüft; nicht unmittelbar erforderliche Teile werden nicht unnötig im ersten sichtbaren Ablauf belastet.", "Bestanden"),
        ("7.2", "DOM-Umfang und Anzahl geladener Ressourcen reduzieren", "Wiederholte beziehungsweise unnötige UI-Strukturen bereinigt und Ressourcenaufkommen im Portal kontrolliert.", "Korrigiert"),
        ("7.3", "Fehlendes Favicon ergänzen", "Favicon ergänzt; Browseranforderung liefert das Symbol statt eines Fehlers.", "Korrigiert"),
        ("7.4", "Automatisierte Tests für zentrale Portalabläufe aufbauen", "Playwright-Suite für Desktop und Mobil eingerichtet. Acht öffentliche Tests bestanden; acht geschützte Tests wurden wegen fehlender Testkonto-Secrets übersprungen.", "Akzeptiert"),
    ],
    [("Automatisierung", "Playwright-Konfiguration, öffentliche und geschützte Portaltests sowie CI-Ablauf im Repository."), ("Ergebnis", "8 öffentliche Tests bestanden (Desktop und Pixel 7); 8 geschützte Tests übersprungen."), ("Ergänzende Abdeckung", "Geschützte Abläufe wurden in Punkt 8.2 manuell mit Mitarbeiter- und Managerkonto abgenommen."), ("Abhängigkeiten", "npm audit meldete 0 bekannte Schwachstellen.")],
    "Punkt 7 ist mit dokumentierter Einschränkung abgeschlossen. Die öffentliche Testsuite ist grün; die fehlende automatisierte Anmeldung der geschützten Tests bleibt als akzeptiertes Restrisiko bestehen.")

add_section_page(doc, 8, "Abnahme und Produktion",
    "Ziel war die nachvollziehbare Produktionsfreigabe nach Vorschau-, Rollen-, Sicherheits- und Regressionstests einschließlich Rückfallmöglichkeit.",
    [
        ("8.1", "Korrekturen zuerst als Vercel-Vorschau bereitstellen", "Korrigierte Versionen wurden vor der Produktivfreigabe erstellt, geprüft und anschließend in das bestehende Vercel-Projekt übernommen.", "Bestanden"),
        ("8.2", "Gemeinsame Abnahme mit Mitarbeiter- und Manager-Testkonto", "Alle 11 Mitarbeiterbereiche und 9 Managerbereiche sowie gemeinsame Vorgänge und Rollenübergänge geprüft.", "Manuell"),
        ("8.3", "Sicherheits- und Regressionstests vollständig wiederholen", "Build und statische Regression bestanden; 31/31 Tabellen mit RLS; keine anonymen Tabellenrechte; npm audit ohne Fund; Sicherheitsheader aktiv.", "Bestanden"),
        ("8.4", "Produktiv veröffentlichen", "Aktives Vercel-Deployment ist READY; Produktionsadresse liefert HTTP 200; keine Fehlerprotokolle bei Freigabe.", "Bestanden"),
        ("8.5", "Rückfallmöglichkeit dokumentieren", "Vorheriges stabiles Deployment als Rückfallziel erfasst; Umschalt- und Rücknahmeverfahren sowie Datenbankhinweis dokumentiert.", "Bestanden"),
    ],
    [("Produktionsadresse", "https://shiftpilot-two.vercel.app/"), ("Aktives Deployment", "dpl_GBstKr5PaoBcbpXwr67L11APwZnQ – Status READY."), ("Rückfallziel", "dpl_6ShoquoGBvoH2D5fD35wpLMfJLW1 – Status READY."), ("Sicherheitsheader", "CSP, HSTS, nosniff, X-Frame-Options DENY, Referrer-Policy und Permissions-Policy geprüft."), ("Datenbank", "Diese Veröffentlichung enthält keine neue Datenbankmigration; der Rückfall schaltet nur das Frontend-Artefakt um.")],
    "Punkt 8 ist abgeschlossen. Die Version ist produktiv freigegeben; Rückfallziel und Betriebsprüfung sind dokumentiert.")

# Residual risks and rollback appendix
doc.add_page_break()
add_heading(doc, "Akzeptierte Restrisiken und Freigabeauflagen", 1)
add_text(doc, "Die folgenden drei Punkte wurden auf ausdrücklichen Wunsch für diese Freigabe übersprungen. Sie sind keine bestandenen Prüfungen, sondern bewusst akzeptierte Restrisiken.", size=10.5, color=MUTED, after=10)
risks = [
    ("R1", "Schutz vor kompromittierten Passwörtern", "In Supabase Free nicht verfügbar. Aktivierungsversuch im Dashboard ergab einen Pro-Tarif-Hinweis; kein Tarifwechsel durchgeführt.", "Akzeptiert"),
    ("R2", "Weitere Härtung von fünf SECURITY-DEFINER-Funktionen", "Supabase meldet fünf öffentlich ausführbare Funktionen. Rollen-/Owner-Prüfungen sind vorhanden; zusätzliche Härtung wurde für diese Freigabe zurückgestellt.", "Akzeptiert"),
    ("R3", "Secrets für acht geschützte automatisierte End-to-End-Tests", "SF_E2E_EMAIL und SF_E2E_PASSWORD sind nicht hinterlegt. Die betroffenen Abläufe wurden manuell mit beiden Rollen abgenommen.", "Akzeptiert"),
]
add_test_table(doc, risks)
add_status_callout(doc, "Freigabebedingung", "Die Freigabe ist gültig, solange diese Risiken bekannt und organisatorisch akzeptiert bleiben. Bei Tarifwechsel, Sicherheitsumbau oder CI-Härtung sind R1 bis R3 erneut zu bewerten.", PALE_GOLD, GOLD)

add_heading(doc, "Rückfallplan", 1)
add_bullet(doc, "Rückfallziel", "Deployment dpl_6ShoquoGBvoH2D5fD35wpLMfJLW1 (Artefakt: https://shiftpilot-kgvb5i548-secure-match.vercel.app).")
add_bullet(doc, "Auslöser", "Schwerwiegender Betriebsfehler der aktuellen Produktion, insbesondere gestörte Anmeldung, fehlerhafte Datentrennung oder nicht nutzbare Kernbereiche.")
add_bullet(doc, "Vorgehen", "Vorheriges Deployment im verknüpften Vercel-Projekt auf die Produktionsadresse promoten und anschließend beide Rollen prüfen.")
add_bullet(doc, "Kontrollschritte", "HTTP 200, Manager-Anmeldung, Mitarbeiter-Anmeldung, persönliche Datentrennung, Managerportal und Schicht-Marktplatz.")
add_bullet(doc, "Rücknahme", "Nach Fehlerbehebung das freigegebene Deployment dpl_GBstKr5PaoBcbpXwr67L11APwZnQ erneut promoten und dieselben Kontrollen wiederholen.")
add_bullet(doc, "Datenbank", "Kein Datenbank-Rollback erforderlich, da diese Veröffentlichung keine neue Migration enthält.")

add_heading(doc, "Abschließende Freigabeentscheidung", 1)
add_status_callout(doc, "PRODUKTIV FREIGEGEBEN", "Das SchichtFunk Mitarbeiterportal wurde nach den Punkten 1 bis 8 geprüft, korrigiert und bereitgestellt. Die Freigabe erfolgte mit den drei oben dokumentierten Restrisiken. Für eine spätere Freigabe ohne diese Einschränkungen müssen R1 bis R3 geschlossen und die geschützten automatisierten Tests erneut vollständig ausgeführt werden.", PALE_TEAL, GREEN)
add_text(doc, "Dokumentationsstand: 01.09.2026  |  Prüfobjekt: SchichtFunk Mitarbeiterportal  |  Vercel-Projekt: shiftpilot", size=9, color=MUTED, italic=True, before=12, after=0)

# Keep table rows intact where possible and set document metadata.
for table in doc.tables:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)

doc.core_properties.title = "SchichtFunk Test- und Freigabedokumentation"
doc.core_properties.subject = "Vollständige Prüfung des Mitarbeiterportals nach Punkt 1 bis 8"
doc.core_properties.author = "SchichtFunk Projektteam"
doc.core_properties.keywords = "SchichtFunk, Mitarbeiterportal, Test, Freigabe, Vercel, Supabase"
doc.save(OUT)
print(OUT)
