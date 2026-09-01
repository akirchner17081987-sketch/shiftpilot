from __future__ import annotations

from datetime import date
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "documentation"
LOGO = OUT / "schichtfunk-logo.png"
BANNER = OUT / "schichtfunk-dokumentation-banner.png"
ARCH = OUT / "schichtfunk-architektur.png"
DOCX = OUT / "SchichtFunk_Projektdokumentation_2026-08-31.docx"

# compact_reference_guide tokens + named SchichtFunk brand override
FONT = "Calibri"
INK = "163047"
MUTED = "5F7386"
TEAL = "159E92"
TEAL_DARK = "0B6D68"
TEAL_LIGHT = "DDF6F2"
BLUE = "2E74B5"
BLUE_LIGHT = "E8EEF5"
NAVY = "071723"
WHITE = "FFFFFF"
GRAY = "F2F4F7"
RULE = "CAD6DF"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def font(path: str, size: int):
    p = Path("C:/Windows/Fonts") / path
    return ImageFont.truetype(str(p), size) if p.exists() else ImageFont.load_default()


def make_visuals():
    logo = Image.open(LOGO).convert("RGBA")
    w, h = 1800, 620
    canvas = Image.new("RGB", (w, h), f"#{NAVY}")
    px = canvas.load()
    for y in range(h):
        for x in range(w):
            t = (x / w) * .55 + (y / h) * .45
            px[x, y] = (7 + int(8*t), 23 + int(17*t), 35 + int(24*t))
    draw = ImageDraw.Draw(canvas, "RGBA")
    for r, alpha in [(470, 48), (350, 40), (230, 32)]:
        draw.ellipse((w-r-90, -r//2, w+120, r+r//2), outline=(42, 222, 202, alpha), width=4)
    for i in range(12):
        x = 80 + i * 138
        draw.line((x, 535, x+110, 535), fill=(42, 222, 202, 45), width=2)
    target_w = 1280
    scale = target_w / logo.width
    logo = logo.resize((target_w, int(logo.height*scale)), Image.Resampling.LANCZOS)
    canvas.paste(logo, ((w-logo.width)//2, 105), logo)
    draw.text((w//2, 520), "PRODUKT- UND SYSTEMDOKUMENTATION  |  STAND 31.08.2026",
              font=font("segoeuib.ttf", 26), fill=(186, 214, 224, 230), anchor="mm")
    canvas.save(BANNER, quality=96)

    aw, ah = 1700, 600
    img = Image.new("RGB", (aw, ah), "#F7FAFC")
    d = ImageDraw.Draw(img)
    title_f = font("segoeuib.ttf", 32)
    box_f = font("segoeuib.ttf", 25)
    small_f = font("segoeui.ttf", 19)
    d.text((70, 45), "SchichtFunk Systemarchitektur", font=title_f, fill="#163047")
    boxes = [
        (70, 175, 365, 390, "Nutzeroberflächen", "Hauptportal\nMitarbeiterportal\nResponsive Web-App"),
        (475, 175, 770, 390, "Anwendungslogik", "Modulare JavaScript-\nFunktionen, Workflows\nund Prüfregeln"),
        (880, 175, 1175, 390, "Supabase", "Authentifizierung\nRPC-Funktionen\nBenachrichtigungen"),
        (1285, 175, 1580, 390, "PostgreSQL", "Mandantendaten\nRLS und Rollen\nAudit-Protokolle"),
    ]
    for idx, (x1,y1,x2,y2,title,body) in enumerate(boxes):
        d.rounded_rectangle((x1,y1,x2,y2), radius=20, fill="#FFFFFF", outline="#8FB4C4", width=3)
        d.rounded_rectangle((x1,y1,x2,y1+58), radius=20, fill="#0D3A43")
        d.rectangle((x1,y1+35,x2,y1+58), fill="#0D3A43")
        d.text(((x1+x2)//2,y1+29), title, font=box_f, fill="#FFFFFF", anchor="mm")
        d.multiline_text((x1+25,y1+88), body, font=small_f, fill="#425E70", spacing=10)
        if idx < len(boxes)-1:
            d.line((x2+18,282,x2+86,282), fill="#159E92", width=7)
            d.polygon([(x2+86,270),(x2+108,282),(x2+86,294)], fill="#159E92")
    d.text((70, 500), "Betrieb: Vercel  |  Datenhaltung: Supabase/PostgreSQL  |  Sicherheit: RLS, Rollen und serverseitige Prüfungen",
           font=small_f, fill="#5F7386")
    img.save(ARCH, quality=95)


def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total)); tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd"); tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(TABLE_INDENT)); tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            trPr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths)-1)]
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
            tcW.set(qn("w:w"), str(width)); tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, bold=None, color=INK, italic=None, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    return run


def add_num_def(doc, kind):
    numbering = doc.part.numbering_part.element
    ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    abs_id = max(ids, default=0) + 1
    nums = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(nums, default=0) + 1
    abstract = OxmlElement("w:abstractNum"); abstract.set(qn("w:abstractNumId"), str(abs_id))
    multi = OxmlElement("w:multiLevelType"); multi.set(qn("w:val"), "singleLevel"); abstract.append(multi)
    lvl = OxmlElement("w:lvl"); lvl.set(qn("w:ilvl"), "0"); abstract.append(lvl)
    start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); lvl.append(start)
    numFmt = OxmlElement("w:numFmt"); numFmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal"); lvl.append(numFmt)
    lvlText = OxmlElement("w:lvlText"); lvlText.set(qn("w:val"), "•" if kind == "bullet" else "%1."); lvl.append(lvlText)
    suff = OxmlElement("w:suff"); suff.set(qn("w:val"), "tab"); lvl.append(suff)
    pPr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab"); tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), "540"); tabs.append(tab); pPr.append(tabs)
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "540"); ind.set(qn("w:hanging"), "270"); pPr.append(ind); lvl.append(pPr)
    numbering.append(abstract)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id))
    abstract_id = OxmlElement("w:abstractNumId"); abstract_id.set(qn("w:val"), str(abs_id)); num.append(abstract_id)
    numbering.append(num)
    return num_id


def list_item(doc, text, num_id):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    pPr = p._p.get_or_add_pPr(); numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId"); num.set(qn("w:val"), str(num_id))
    numPr.extend([ilvl, num]); pPr.append(numPr)
    set_run(p.add_run(text))
    return p


def add_field(paragraph, code):
    run = paragraph.add_run(); begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = code
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t"); txt.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, txt, end])


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); tblHeader = OxmlElement("w:tblHeader"); tblHeader.set(qn("w:val"), "true"); trPr.append(tblHeader)


def table(doc, headers, rows, widths, font_size=9.2):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    set_table_geometry(t, widths)
    set_repeat_table_header(t.rows[0])
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]; shade(cell, TEAL_DARK)
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run(p.add_run(h), size=9.2, bold=True, color=WHITE)
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            if ridx % 2: shade(cells[i], "F7FAFC")
            p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.12
            set_run(p.add_run(str(value)), size=font_size, color=INK)
    set_table_geometry(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t


def heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True)
        set_run(p.add_run(text[len(bold_lead):]))
    else:
        set_run(p.add_run(text))
    return p


def callout(doc, title, text, fill=TEAL_LIGHT):
    t = doc.add_table(rows=1, cols=1); set_table_geometry(t, [TABLE_WIDTH])
    cell = t.cell(0,0); shade(cell, fill); set_cell_margins(cell, 160, 180, 160, 180)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(title), size=10.5, bold=True, color=TEAL_DARK)
    p2 = cell.add_paragraph(); p2.paragraph_format.space_after = Pt(0); p2.paragraph_format.line_spacing = 1.2
    set_run(p2.add_run(text), size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_after = Pt(0)


def add_feature(doc, title, purpose, functions, benefit, bullet_id, level=3):
    heading(doc, title, level)
    para(doc, purpose)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run("Kernfunktionen"), size=10.5, bold=True, color=TEAL_DARK)
    for item in functions: list_item(doc, item, bullet_id)
    para(doc, f"Nutzen: {benefit}", bold_lead="Nutzen:")


def configure_document(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT; normal._element.rPr.rFonts.set(qn("w:ascii"), FONT); normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11); normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
    for level, size, before, after in [(1,16,18,10),(2,13,14,7),(3,12,10,5)]:
        st = styles[f"Heading {level}"]; st.font.name = FONT; st._element.rPr.rFonts.set(qn("w:ascii"), FONT); st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = RGBColor.from_string(TEAL_DARK if level < 3 else INK)
        st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after); st.paragraph_format.keep_with_next = True
    doc.settings.odd_and_even_pages_header_footer = True
    for current_header in (sec.header, sec.even_page_header):
        hp = current_header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run(hp.add_run("SchichtFunk"), size=9, bold=True, color=TEAL_DARK)
        set_run(hp.add_run("  |  Produkt- und Systemdokumentation"), size=9, color=MUTED)
        pPr = hp._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "8"); bottom.set(qn("w:space"), "5"); bottom.set(qn("w:color"), TEAL)
        pBdr.append(bottom); pPr.append(pBdr)
    for current_footer in (sec.footer, sec.even_page_footer):
        fp = current_footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_run(fp.add_run("Stand 31.08.2026  |  Seite "), size=8.5, color=MUTED); add_field(fp, "PAGE")


def build():
    OUT.mkdir(parents=True, exist_ok=True)
    make_visuals()
    doc = Document(); configure_document(doc)
    bullet_id = add_num_def(doc, "bullet")

    # Cover
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(24)
    pic = p.add_run().add_picture(str(BANNER), width=Inches(6.5))
    pic._inline.docPr.set("descr", "SchichtFunk Firmenlogo und Dokumentationsbanner")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run("Projekthistorie und vollständige Funktionsdokumentation"), size=24, bold=True, color=INK)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(20)
    set_run(p.add_run("Hauptportal, Mitarbeiterportal, Prozesse, Sicherheit und Betrieb"), size=13, color=MUTED)
    table(doc, ["Dokument", "Stand", "Produktstatus"], [["Produkt- und Systemhandbuch", "31. August 2026", "Produktiv auf Vercel"]], [3500, 2300, 3560], 9.5)
    callout(doc, "Zweck des Dokuments", "Diese Dokumentation beschreibt den nachvollziehbaren Entwicklungsweg von ShiftPilot zu SchichtFunk sowie alle aktuell implementierten Funktionsbereiche. Sie dient als internes Referenzhandbuch und als verständliche Produktübersicht für Kunden, Partner und Projektbeteiligte.")
    para(doc, "Dokumentationsbasis: aktueller Quellcode, Datenbankmigrationen, automatisierte Regressionstests und 448 nachvollziehbare Versionsstände im Repository. Aktueller Release: a950f32.")

    page_break(doc)
    heading(doc, "Dokumentstruktur", 1)
    structure_id = add_num_def(doc, "decimal")
    for item in [
        "Produktprofil und Leitidee",
        "Projekthistorie und Meilensteine",
        "Systemarchitektur und Datenfluss",
        "Rollen, Berechtigungen und Sicherheit",
        "Öffentliche Hauptseite",
        "Hauptportal für Planung und Verwaltung",
        "Mitarbeiterportal",
        "Übergreifende Workflows und Automatisierung",
        "Qualitätssicherung, Betrieb und Veröffentlichung",
        "Technischer Modul- und Datenbankanhang",
    ]: list_item(doc, item, structure_id)
    callout(doc, "Lesetipp", "Die Kapitel 6 und 7 sind die eigentliche Funktionsreferenz. Kapitel 2 erklärt die Entwicklungsgeschichte; Kapitel 3, 4 und 10 richten sich besonders an technische und organisatorische Verantwortliche.", BLUE_LIGHT)

    page_break(doc)
    heading(doc, "1. Produktprofil und Leitidee", 1)
    para(doc, "SchichtFunk ist eine webbasierte Plattform für Dienstplanung, Personaleinsatz, Arbeitszeit, Abwesenheiten und operative Reaktion auf kurzfristige Änderungen. Der Produktname verbindet den Kernprozess Schichtplanung mit dem Gedanken schneller, nachvollziehbarer Kommunikation.")
    callout(doc, "Markenversprechen", "Klar geplant. Stark besetzt. - SchichtFunk schafft einen gemeinsamen, nachvollziehbaren Arbeitsbereich für Disposition, Führungskräfte und Mitarbeitende.")
    heading(doc, "1.1 Zielgruppen", 2)
    for item in [
        "Unternehmensinhaber und Administratoren, die Organisation, Rechte und Regeln verwalten.",
        "Disposition und Planung, die Dienstpläne erstellen, prüfen, veröffentlichen und Störfälle lösen.",
        "Führungskräfte, die Personal, Zeiten, Abwesenheiten und Auswertungen kontrollieren.",
        "Mitarbeitende, die ihre persönlichen Schichten, Anträge, Zeiten und Konten einsehen und bearbeiten.",
    ]: list_item(doc, item, bullet_id)
    heading(doc, "1.2 Produktprinzipien", 2)
    table(doc, ["Prinzip", "Bedeutung"], [
        ["Eine Datenbasis", "Dienstplan, Personal, Zeiten, Abwesenheiten und Workflows greifen auf gemeinsame Unternehmensdaten zu."],
        ["Kontrollierte Automatik", "Vorschläge werden automatisch erzeugt, aber nachvollziehbar geprüft und bewusst übernommen."],
        ["Rechte vor Oberfläche", "Berechtigungen werden nicht nur visuell, sondern serverseitig und in der Datenbank durchgesetzt."],
        ["Nachvollziehbarkeit", "Änderungen, Freigaben und sicherheitsrelevante Vorgänge werden protokolliert."],
        ["Zwei passende Portale", "Verwaltung und Mitarbeitende erhalten getrennte, auf ihre Aufgaben zugeschnittene Arbeitsbereiche."],
    ], [2100, 7260], 9.4)

    heading(doc, "2. Projekthistorie und Meilensteine", 1)
    para(doc, "Die Entwicklung erfolgte in einer sehr kompakten, iterativen Aufbauphase vom 25. bis 31. August 2026. Insgesamt sind 448 Versionsstände dokumentiert. Die Historie zeigt den Übergang von einem lokalen Planungsprototyp zu einer produktiv betriebenen, mandantenfähigen Anwendung mit Supabase-Backend und zwei Rollenportalen.")
    table(doc, ["Zeitraum", "Schwerpunkt", "Wesentliche Ergebnisse"], [
        ["25.08.2026", "Grundprodukt", "Initiales ShiftPilot-Projekt, Dienstplan, Drag-and-drop, SOLL/IST, Mehrfachbesetzung, erste Dashboards und Einstellungen."],
        ["26.08.2026", "Funktionsausbau und UX", "Landingpage, Vorlagenverwaltung, Zeiterfassung, Auto-Planung, Abwesenheiten, Auswertungen sowie umfassende Bedien- und Lesbarkeitsverbesserungen."],
        ["27.08.2026", "Regeln und Rebranding", "OT-Wochenend-/Feiertagsregeln, dynamischer Mitarbeiterpool, zentrale Konfliktprüfung und Umbenennung von ShiftPilot zu SchichtFunk."],
        ["28.08.2026", "Cloud und Compliance", "Supabase Auth, PostgreSQL-Synchronisierung, Compliance-Workflow, Legacy-Migration, Rollen und erstes Mitarbeiterportal."],
        ["29.08.2026", "Mitarbeiterprozesse", "Passwort-Reset, Schichtänderungsantworten, Abwesenheitsanträge und sicherer Dienstplan-Reset."],
        ["30.08.2026", "Unternehmensfunktionen", "Benachrichtigungen, Schichttausch, Cloud-Zeiterfassung, Stundenkonto, Feiertage, Exporte, Monatsabschluss, Personalakte, Fristen und neue Wochen-/Monatsansichten."],
        ["31.08.2026", "Reifegrad und Innovation", "Vollständige Bereiche, Marktplatz, Regressionstests, Sicherheits-Härtung, Mobiloptimierung, Einsatzbereitschafts-Ampel, Störfall-Autopilot und neues Mitarbeiterportal."],
    ], [1500, 2500, 5360], 8.8)
    heading(doc, "2.1 Zentrale Entwicklungslinien", 2)
    for item in [
        "Vom lokalen Planungsmodell zur cloudbasierten, mandantengetrennten Datenhaltung.",
        "Vom einzelnen Dienstplan zur verknüpften Plattform aus Planung, Personal, Zeit und Abwesenheit.",
        "Vom Administrationsbereich zum getrennten, persönlichen Mitarbeiterportal.",
        "Von manueller Planung zu kontrollierter Auto-Planung, Einsatzbereitschafts-Ampel und Störfall-Autopilot.",
        "Von reinem UI-Schutz zu Datenbankregeln, RLS, abgesicherten RPC-Funktionen und Regressionstests.",
    ]: list_item(doc, item, bullet_id)
    heading(doc, "2.2 Markenentwicklung", 2)
    para(doc, "Das Projekt startete unter dem Namen ShiftPilot. Am 27. August 2026 wurde das Rebranding zu SchichtFunk umgesetzt. Logo, Texte, Landingpage, Portale und Exporte wurden anschließend schrittweise vereinheitlicht. Der Claim 'Klar geplant. Stark besetzt.' wurde zum verbindenden Leitmotiv.")

    heading(doc, "3. Systemarchitektur und Datenfluss", 1)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8)
    pic = p.add_run().add_picture(str(ARCH), width=Inches(6.5)); pic._inline.docPr.set("descr", "Diagramm der SchichtFunk Systemarchitektur")
    para(doc, "Die Anwendung ist modular aufgebaut. Die Benutzeroberfläche wird als responsive Web-App ausgeliefert; Supabase übernimmt Authentifizierung, Datenzugriff und serverseitige Funktionen. PostgreSQL erzwingt Mandantentrennung und Fachregeln. Vercel stellt die produktive Anwendung bereit.")
    heading(doc, "3.1 Technologiestapel", 2)
    table(doc, ["Ebene", "Technologie", "Aufgabe"], [
        ["Frontend", "HTML, CSS, JavaScript", "Landingpage, Hauptportal, Mitarbeiterportal und interaktive Workflows."],
        ["Modulstruktur", "Dynamisch geladene Assets", "Fachbereiche bleiben getrennt erweiterbar und cachebar."],
        ["Authentifizierung", "Supabase Auth", "Anmeldung, Sitzungen, Einladungen und Passwort-Reset."],
        ["Datenbank", "PostgreSQL / Supabase", "Unternehmens-, Personal-, Planungs-, Zeit- und Auditdaten."],
        ["Sicherheit", "RLS und geschützte RPCs", "Mandantentrennung, Rollenprüfung und atomare Fachprozesse."],
        ["Hosting", "Vercel", "Produktionsbereitstellung, Versionierung und weltweite Auslieferung."],
        ["Qualität", "Node Regressionstests", "Prüfung kritischer Verlinkungen, Module und Sicherheitsmuster."],
    ], [1650, 2450, 5260], 9)
    heading(doc, "3.2 Grundlegender Datenfluss", 2)
    dataflow_id = add_num_def(doc, "decimal")
    for text in [
        "Die Nutzerin oder der Nutzer meldet sich über Supabase Auth an.",
        "Die Anwendung ermittelt Unternehmen, Rolle und gegebenenfalls den zugehörigen Mitarbeiterdatensatz.",
        "Das Portal lädt ausschließlich die für Rolle und Unternehmen zulässigen Daten.",
        "Schreibvorgänge laufen über abgesicherte Datenbankfunktionen oder RLS-geschützte Tabellen.",
        "Fachregeln prüfen Überschneidungen, Ruhezeiten, Freigaben, Status und Veröffentlichungsstand.",
        "Benachrichtigungen und Audit-Ereignisse machen Ergebnisse und Änderungen nachvollziehbar.",
    ]: list_item(doc, text, dataflow_id)

    heading(doc, "4. Rollen, Berechtigungen und Sicherheit", 1)
    para(doc, "SchichtFunk trennt Unternehmensverwaltung, operative Planung und persönliche Mitarbeiterdaten. Die sichtbare Oberfläche folgt der Rolle; entscheidend ist jedoch die serverseitige Prüfung in Supabase/PostgreSQL.")
    table(doc, ["Rolle", "Typische Befugnisse", "Schutzgrenzen"], [
        ["Inhaber (OWNER)", "Vollständige Unternehmens- und Benutzerverwaltung.", "Inhaberrolle ist besonders geschützt und nicht beliebig übertragbar."],
        ["Administrator (ADMIN)", "Benutzer, Einstellungen, Personal, Planung und Auswertungen verwalten.", "Keine unkontrollierte Selbsterhöhung oder Änderung geschützter Rollen."],
        ["Disposition (DISPATCHER)", "Dienstplan, operative Änderungen, Marktplatz und Störfälle bearbeiten.", "Kein vollständiger Zugriff auf geschützte Unternehmensverwaltung."],
        ["Planung (PLANNER)", "Planungs- und prüfbezogene Aufgaben ausführen.", "Berechtigungen werden pro Datenbankfunktion kontrolliert."],
        ["Betrachter (VIEWER)", "Zulässige Unternehmensinformationen lesen.", "Keine administrativen Schreibrechte."],
        ["Mitarbeiter", "Nur eigene Schichten, Anträge, Zeiten, Konten und Angebote bearbeiten.", "Keine fremden Personaldaten und keine Verwaltungsfunktionen."],
    ], [1550, 4240, 3570], 8.8)
    heading(doc, "4.1 Sicherheitsmechanismen", 2)
    for item in [
        "Row Level Security (RLS) begrenzt Tabellenzugriffe auf das aktive Unternehmen und die eigene Identität.",
        "Privilegierte Funktionen prüfen auth.uid(), Mitgliedschaft, Rollenstatus und Fachbedingungen erneut.",
        "Direkte Tabellenrechte werden auf das notwendige Minimum reduziert; Ausführung erfolgt über definierte Funktionen.",
        "Atomare Transaktionen und Zeilensperren verhindern doppelte oder konkurrierende Übernahmen.",
        "Audit-Ereignisse protokollieren relevante Status- und Zuordnungsänderungen.",
        "Veröffentlichte Dienstpläne und persönliche Portalansichten besitzen gesonderte Sichtbarkeitsregeln.",
        "Eine behobene Rollen-Eskalation wird durch Regressionstests dauerhaft abgesichert.",
    ]: list_item(doc, item, bullet_id)
    callout(doc, "Wichtiger Hinweis", "SchichtFunk unterstützt die organisatorische und technische Einhaltung hinterlegter Regeln. Die Anwendung ersetzt keine individuelle arbeitsrechtliche Prüfung oder Rechtsberatung.", "FFF6DA")

    heading(doc, "5. Öffentliche Hauptseite", 1)
    para(doc, "Die öffentliche Hauptseite präsentiert SchichtFunk als moderne Dienstplanungsplattform und führt Interessierte gezielt in die Anwendung. Sie ist vom geschützten Arbeitsbereich getrennt.")
    add_feature(doc, "5.1 Start- und Markenbereich", "Der Hero-Bereich führt Marke, Nutzenversprechen und Einstieg zusammen.", ["SchichtFunk Logo und Claim", "Direkte Anmeldung und Dashboard-Einstieg", "Produktvorschau im dunklen SchichtFunk-Design", "Kurze Nutzenargumente zu Einfachheit, Zeitersparnis, Fehlervermeidung und Übersicht"], "Besucher verstehen Produktzweck und Mehrwert unmittelbar.", bullet_id, level=2)
    add_feature(doc, "5.2 Funktionsübersicht", "Sechs Funktionskarten erklären die zentralen Produktbereiche.", ["Intelligente Dienstplanung", "Mitarbeiterverwaltung", "Arbeitszeit und Stundenkonto", "Abwesenheiten", "Auto-Planung und Prüfungen", "Auswertungen"], "Die Leistungsbreite ist ohne Anmeldung nachvollziehbar.", bullet_id, level=2)
    add_feature(doc, "5.3 Vorteile und Lösungsweg", "Die Seite verbindet Nutzenargumente mit einem vierstufigen Prozessmodell.", ["Team und Regeln einrichten", "Dienstplan erstellen", "Konflikte automatisch prüfen", "Veröffentlichen, zusammenarbeiten und auswerten", "Mitarbeiterportal als eigener Vorteil"], "Interessierte erkennen, wie die Bereiche zusammenspielen.", bullet_id, level=2)
    add_feature(doc, "5.4 Kontakt und Abschluss", "Kontaktbereich und Footer schließen die Produktdarstellung ab.", ["Handlungsaufforderungen zur Anwendung", "Navigation zu Funktionen, Vorteilen und Lösung", "Markenkonsistenter Abschlussbereich"], "Der Weg vom Produktinteresse zum geschützten Arbeitsbereich bleibt klar.", bullet_id, level=2)

    heading(doc, "6. Hauptportal für Planung und Verwaltung", 1)
    para(doc, "Das Hauptportal ist der operative Arbeitsbereich für Verwaltung, Disposition und Planung. Eine linke Seitenleiste steuert die Bereiche; Kopfzeile, globale Suche, Cloud-Status, Benachrichtigungen, Hilfe und Schnellaktionen bleiben verfügbar.")
    heading(doc, "6.1 Bereichsübersicht", 2)
    table(doc, ["Bereich", "Aufgabe", "Kernfunktionen"], [
        ["Übersicht", "Aktuelle Lage erfassen", "Kennzahlen, Handlungsbedarf, Abwesenheiten, Auslastung, Schichtmix, Schnellaktionen"],
        ["Dienstplan", "Schichten planen und veröffentlichen", "Woche/Monat, Vorlagen, SOLL/IST, Zuweisung, Compliance, Verlauf"],
        ["Mitarbeiter", "Personal verwalten", "Stammdaten, Freigaben, Zugänge, Personalakte, Fristen"],
        ["Zeiterfassung", "SOLL und IST steuern", "Ist-Zeiten, Abweichungen, Korrekturen, Bestätigungen, Konten"],
        ["Abwesenheiten", "Verfügbarkeit organisieren", "Anträge, Entscheidungen, Kalender, Konflikte"],
        ["Auto-Planung", "Offene Positionen vorschlagen", "Regeln, Analyse, Vorschläge, ungelöste Positionen"],
        ["Störfall-Autopilot", "Akute Ausfälle lösen", "Kandidatenranking, Anfragen, erste Zusage, Audit"],
        ["Schicht-Marktplatz", "Schichten kontrolliert vermitteln", "Angebote, Übernahmen, Fachprüfung, Freigabe"],
        ["Auswertungen", "Betrieb analysieren", "Besetzung, Arbeitszeit, Abwesenheit, Auffälligkeiten, Export"],
        ["Einstellungen", "Regeln und Organisation pflegen", "Schichtzeiten, SOLL-Stärke, Benutzer, Rollen und Kontoeinstellungen"],
    ], [1750, 2740, 4870], 8.5)

    add_feature(doc, "6.2 Planungs-Dashboard", "Die Übersicht bündelt die wichtigsten Informationen der aktuellen Woche und macht Handlungsbedarf direkt sichtbar.", ["Planungsfortschritt und Besetzungsquote", "Offene Positionen und heutiger Dienst", "Abwesenheiten und aktive Mitarbeiter", "Mitarbeiter-Auslastung und Schichtverteilung", "Direkte Schnellnavigation in Dienstplan, Mitarbeiter, Abwesenheiten und Auto-Planung"], "Führungskräfte erkennen ohne Detailnavigation, wo Eingriffe nötig sind.", bullet_id)
    add_feature(doc, "6.3 Dienstplan", "Der Dienstplan bildet den Kern der operativen Personaleinsatzplanung.", ["Wochen- und Monatsansicht", "Schichtbibliothek mit verwaltbaren Vorlagen", "Drag-and-drop und direkte Mitarbeiterzuweisung", "Dynamischer Mitarbeiterpool mit Suche, Freigaben und Verfügbarkeit", "SOLL/IST-Besetzung und offene Positionen", "Bearbeiten, Schnelllöschen, Verlauf und kompletter Reset", "Veröffentlichungsstatus und atomare Wochenfreigabe", "Konflikt-, Plausibilitäts- und Complianceprüfung", "Einsatzbereitschafts-Ampel für Besetzung, Regeln und Bestätigungen"], "Planung, Prüfung und Veröffentlichung bleiben in einem durchgängigen Ablauf.", bullet_id)
    add_feature(doc, "6.4 Mitarbeiterverwaltung", "Der Personalbereich bündelt alle für Planung und Verwaltung benötigten Mitarbeiterinformationen.", ["Stammdaten, Personalnummer, Beschäftigung und Wochenstunden", "Kontaktdaten und bevorzugter Kontaktweg", "Schichtfreigaben und Verfügbarkeiten", "Status, Suche, Filter und responsive Darstellung", "Mitarbeiterzugang mit Einladung und Sperrstatus", "Personalakte mit internen Daten, Dokumentenhinweisen und Notfallkontakt", "Fristen-Dashboard und Erinnerungsnavigation", "Schutz vor doppelten Personalnummern und selbstheilender Cloud-Synchronisierung"], "Personaldaten bleiben zentral, planungsrelevant und nachvollziehbar.", bullet_id)
    add_feature(doc, "6.5 Zeiterfassung und Stundenkonto", "Der Arbeitszeitbereich verbindet geplante Schichten mit gemeldeten und bestätigten Ist-Zeiten.", ["SOLL-/IST-Vergleich, Pausen und Abweichungen", "Mitarbeiter melden Ist-Zeiten erst nach Schichtende", "Prüfung, Bestätigung und Korrekturanforderung", "Monats- und kumuliertes Stundenkonto", "Feiertagsgutschriften nach Bundesland", "Monatsabschluss mit Sperrfunktion und Wiederöffnung", "Gemeinsame Monatssteuerung für Konto und Lohnvorschau"], "Arbeitszeit wird von der Planung bis zum Monatsabschluss durchgängig abgebildet.", bullet_id)
    add_feature(doc, "6.6 Abwesenheitsmanagement", "Urlaub, Krankheit, Frei, Fortbildung, Sperrzeiten und weitere Arten werden zentral verwaltet.", ["Zeitraum, Art, Status und Hinweis", "Mitarbeitersuche und eigenständige Mitarbeiteranträge", "Genehmigung oder Ablehnung durch Berechtigte", "Wochenkalender und Konfliktauswirkungen", "Direkte Berücksichtigung in Dienstplan und Auto-Planung", "Benachrichtigungen und sichere Modalabläufe"], "Abwesenheiten wirken frühzeitig in die Besetzungsentscheidung ein.", bullet_id)
    add_feature(doc, "6.7 Auto-Planung", "Die Auto-Planung analysiert offene Positionen und erzeugt kontrollierbare Besetzungsvorschläge.", ["Aktivstatus, Schichtfreigabe und Abwesenheit", "Vermeidung von Doppelbelegung und Zeitüberschneidung", "Berücksichtigung von Ruhezeit und Wochenstunden", "Faire Verteilung der Auslastung", "Vorschau, Übernahme oder Verwerfen", "Getrennte Anzeige nicht automatisch lösbarer Positionen"], "Routineplanung wird beschleunigt, ohne die Entscheidungshoheit abzugeben.", bullet_id)
    add_feature(doc, "6.8 Schicht-Marktplatz", "Mitarbeitende können eigene veröffentlichte Schichten anbieten und passende Angebote übernehmen.", ["Eigene Schicht in den Marktplatz stellen", "Passende Angebote nach Freigaben und Regeln anzeigen", "Übernahme zunächst zur Fach- oder Dispositionsprüfung einreichen", "Manager-Dashboard für offene und abgeschlossene Vorgänge", "Zurückziehen, Freigeben und vollständige Statushistorie"], "Schichten werden teamorientiert vermittelt, ohne den Dienstplan unkontrolliert zu verändern.", bullet_id)
    add_feature(doc, "6.9 Störfall-Autopilot", "Der Autopilot unterstützt bei akutem Ausfall einer bereits besetzten und veröffentlichten Schicht.", ["Störfall mit Art und internem Hinweis erfassen", "Geeignete Ersatzkräfte nach Freigabe, Ruhezeit, Verfügbarkeit und Auslastung priorisieren", "Top-3 oder einzelne Kandidaten anfragen", "Zusage oder Ablehnung im Mitarbeiterportal", "Erste gültige Zusage übernimmt die Schicht atomar", "Andere Angebote werden automatisch geschlossen", "Manager-Benachrichtigung und Audit-Protokoll"], "Akute Lücken werden schneller und regelkonform geschlossen.", bullet_id)
    add_feature(doc, "6.10 Auswertungen und Exporte", "Der Berichtbereich verdichtet operative Daten für Kontrolle und Weiterverarbeitung.", ["SOLL-/IST-Besetzung", "Arbeitsstunden und Auslastung", "Abwesenheiten und Schichtverteilung", "Auffälligkeiten", "Monatsauswertungen als Excel und PDF", "Firmenbezogenes PDF-Branding, Tabellenlayout und Saldo-Farben"], "Entscheidungen, Nachweise und Abrechnungsvorbereitung werden unterstützt.", bullet_id)
    add_feature(doc, "6.11 Einstellungen, Benutzer und Rollen", "Zentrale Einstellungen definieren die organisatorische Grundlage.", ["Allgemeine Schichtzeiten", "Globale und tageweise SOLL-Stärke", "OT-Regel für Wochenende und Feiertage", "Unternehmensmitglieder einladen, deaktivieren und Rollen ändern", "Inhaber- und Administratorschutz", "Konto- und Unternehmenskontext"], "Regeln und Zugänge werden an einer kontrollierten Stelle gepflegt.", bullet_id)
    add_feature(doc, "6.12 Benachrichtigungen, Hilfe und Navigation", "Übergreifende Bedienelemente verbinden die Fachbereiche.", ["Benachrichtigungszentrale mit Ungelesen-Zähler", "Direkte Navigation zum betroffenen Vorgang", "Löschen und Gelesen-Markierung", "Kontextbezogenes Hilfe-Center", "Globale Suche und Cloud-Status", "Stabiler Ansichts- und Scrollzustand bei Navigation und Seitenwechsel"], "Wichtige Ereignisse führen ohne Umwege zur richtigen Aufgabe.", bullet_id)

    heading(doc, "7. Mitarbeiterportal", 1)
    para(doc, "Das Mitarbeiterportal ist ein vollständig getrennter persönlicher Arbeitsbereich. Es zeigt ausschließlich die Daten des angemeldeten Mitarbeiters. Ein Unternehmens-Header benennt den Arbeitgeber; eine feste Navigation am linken Rand führt zu klar getrennten Kategorien. Pro Auswahl wird nur der benötigte Bereich angezeigt.")
    heading(doc, "7.1 Aufbau und Navigation", 2)
    for item in [
        "Unternehmens-Header mit Firmenname, sicherem Verbindungshinweis, Benachrichtigungen und Abmeldung.",
        "Linke Kategorienavigation analog zum Hauptportal.",
        "Startübersicht mit Kennzahlen und Schnellzugriff auf alle persönlichen Bereiche.",
        "Einzelansichten statt einer langen, unkoordinierten Kartenseite.",
        "Responsive Mobilnavigation unterhalb des Headers.",
        "Erhalt des zuletzt gewählten Bereichs innerhalb der Sitzung.",
    ]: list_item(doc, item, bullet_id)
    heading(doc, "7.2 Bereiche des Mitarbeiterportals", 2)
    table(doc, ["Kategorie", "Funktion", "Aktionen und Informationen"], [
        ["Übersicht", "Persönlichen Status erfassen", "Kommende Schichten, nächste Schicht, Stunden der nächsten sieben Tage und Schnellzugriffe."],
        ["Ersatzanfragen", "Dringende Anfragen beantworten", "Geprüfte Ersatzschicht ansehen, verbindlich übernehmen oder ablehnen."],
        ["Schicht-Marktplatz", "Schichten anbieten oder übernehmen", "Eigene Angebote, verfügbare Schichten, Eignung, Einreichung und Rückzug."],
        ["Meine Schichten", "Persönlichen Dienstplan sehen", "Nur veröffentlichte eigene Schichten, Zeiten, Codes und Status."],
        ["Schichtänderungen", "Änderungen beantworten", "Anfragen, alte/neue Werte, Bestätigung oder Ablehnung."],
        ["Schichttausch", "Tauschvorgänge verwalten", "Tausch anstoßen, Kollegenantwort und Managerstatus verfolgen."],
        ["Arbeitszeit", "Ist-Zeit melden", "Tatsächlichen Beginn, Ende und Pause erfassen; Korrekturen bearbeiten."],
        ["Abwesenheiten", "Persönliche Anträge", "Antrag stellen, Zeitraum/Art dokumentieren und Entscheidung sehen."],
        ["Stundenkonto", "Monatswerte prüfen", "SOLL, bestätigte Arbeit, Gutschriften, Monatssaldo und Gesamtkonto."],
        ["Lohnvorschau", "Persönliche Vorschau", "Monatsbezogene private Berechnung auf Basis verfügbarer Zeitdaten."],
        ["Mein Profil", "Eigene Stammdaten", "Personalnummer, Beschäftigung, Wochenstunden, Kontakt und Arbeitszeitmodell."],
    ], [1800, 2700, 4860], 8.4)
    heading(doc, "7.3 Persönliche Datenschutzgrenze", 2)
    callout(doc, "Prinzip der Datensparsamkeit", "Das Mitarbeiterportal zeigt keine fremden Mitarbeiterdaten und keine internen Verwaltungsinformationen. Der Zugriff ist an das Auth-Konto und den zugeordneten Mitarbeiterdatensatz gebunden.")
    for item in [
        "Nur eigene veröffentlichte Schichten werden angezeigt.",
        "Angebote, Anträge, Zeiten und Konten sind auf die eigene Mitarbeiter-ID begrenzt.",
        "Interne Personalaktennotizen bleiben ausschließlich im berechtigten Hauptportal.",
        "Deaktivierte oder nicht zugeordnete Mitarbeiterkonten werden blockiert.",
    ]: list_item(doc, item, bullet_id)

    heading(doc, "8. Übergreifende Workflows und Automatisierung", 1)
    heading(doc, "8.1 Dienstplan von Entwurf bis Veröffentlichung", 2)
    publish_id = add_num_def(doc, "decimal")
    for item in ["Vorlagen und SOLL-Stärke festlegen.", "Mitarbeiter manuell oder über Vorschläge zuweisen.", "Konflikte, Abwesenheiten, Freigaben und Compliance prüfen.", "Offene Befunde bearbeiten oder begründet bestätigen.", "Woche atomar veröffentlichen.", "Mitarbeitende sehen ausschließlich den freigegebenen Stand."]: list_item(doc, item, publish_id)
    heading(doc, "8.2 Schichtänderung", 2)
    change_id = add_num_def(doc, "decimal")
    for item in ["Änderung wird angelegt und gegen No-Op sowie Regeln geprüft.", "Falls erforderlich, erhält der Mitarbeiter eine Anfrage.", "Bestätigung oder Ablehnung wird protokolliert.", "Erst nach vollständiger Freigabe wird die Änderung angewendet."]: list_item(doc, item, change_id)
    heading(doc, "8.3 Abwesenheit", 2)
    absence_id = add_num_def(doc, "decimal")
    for item in ["Mitarbeiter oder Verwaltung erfasst den Antrag.", "Berechtigte Stelle entscheidet.", "Genehmigte Zeiträume sperren unzulässige Planung und beeinflussen Vorschläge.", "Konflikte und Benachrichtigungen verknüpfen Abwesenheit und Dienstplan."]: list_item(doc, item, absence_id)
    heading(doc, "8.4 Marktplatz und Störfall", 2)
    para(doc, "Beide Prozesse vermitteln Schichten, unterscheiden sich aber im Auslöser: Der Marktplatz ist ein geplanter, freiwilliger Austausch; der Störfall-Autopilot reagiert auf einen akuten Ausfall. In beiden Fällen verhindern serverseitige Regeln eine unzulässige direkte Planänderung.")
    table(doc, ["Merkmal", "Schicht-Marktplatz", "Störfall-Autopilot"], [
        ["Auslöser", "Mitarbeiter bietet eigene Schicht an", "Disposition meldet akuten Ausfall"],
        ["Auswahl", "Passende Angebote im Portal", "Priorisiertes Kandidatenranking"],
        ["Entscheidung", "Übernahme plus Managerprüfung", "Erste gültige Zusage gewinnt"],
        ["Planänderung", "Kontrollierter Freigabeprozess", "Atomare automatische Übernahme"],
        ["Nachweis", "Status und Audit-Ereignisse", "Status, Benachrichtigungen und Audit"],
    ], [1800, 3780, 3780], 9)
    heading(doc, "8.5 Einsatzbereitschafts-Ampel", 2)
    para(doc, "Die Ampel fasst Besetzung, Freigaben, Ruhezeiten, Veröffentlichungsstatus und Mitarbeiterbestätigungen je Schicht zusammen. Grün bedeutet nach hinterlegten Daten einsatzbereit; Gelb verlangt Prüfung; Rot kennzeichnet kritischen Handlungsbedarf. Sie ist eine Entscheidungshilfe, keine automatische Rechtsfreigabe.")

    heading(doc, "9. Qualitätssicherung, Betrieb und Veröffentlichung", 1)
    heading(doc, "9.1 Automatisierte Regression", 2)
    for item in [
        "Prüfung aller lokal referenzierten Skripte und Styles.",
        "Syntaxprüfung der JavaScript-Module.",
        "Navigationstest für statische und dynamische Bereiche.",
        "Absicherung von Mitarbeiter-E-Mail, Teamleiter-Freigaben und Marktplatznavigation.",
        "Tests gegen Rollen-Eskalation und überprivilegierte Datenbankfunktionen.",
        "Prüfung von Mobilregeln, Einsatzbereitschaft und Schichtbestätigung.",
        "Störfalltests für Mandantentrennung, Sperrreihenfolge und erste Zusage.",
        "Layouttest für Unternehmens-Header und linke Mitarbeiterportal-Navigation.",
    ]: list_item(doc, item, bullet_id)
    heading(doc, "9.2 Durchgeführte Ende-zu-Ende-Prüfungen", 2)
    for item in [
        "Navigation des angemeldeten Hauptportals und Statuswiederherstellung.",
        "Schicht-Marktplatz mit zwei Teamleiter-Testmitarbeitern.",
        "Datenbankseitiger Störfallablauf mit Kandidatenranking, Angebot, Annahme und Rollback.",
        "Gleichzeitige Zusagen: erste Annahme erfolgreich, zweite sicher blockiert.",
        "RLS-Prüfung mit nicht zugehörigem Benutzer: keine sichtbaren Störfalldaten.",
        "Mobile Darstellung und Bedienbarkeit der zentralen Arbeitsbereiche.",
    ]: list_item(doc, item, bullet_id)
    heading(doc, "9.3 Produktionsbetrieb", 2)
    table(doc, ["Aspekt", "Umsetzung"], [
        ["Produktionsadresse", "https://shiftpilot-two.vercel.app/"],
        ["Deployment", "Automatisch über Git-Push auf den Hauptbranch und Vercel."],
        ["Datenbank", "Supabase-Projekt mit PostgreSQL, Auth und produktiven Migrationen."],
        ["Cache-Steuerung", "Versionierte Asset-URLs erzwingen die Auslieferung aktualisierter Module."],
        ["Fehlerkontrolle", "Laufzeitfehler werden nach Veröffentlichungen kontrolliert."],
        ["Aktueller dokumentierter Release", "a950f32 - Mitarbeiterportal-Navigation links."],
    ], [2500, 6860], 9.2)
    callout(doc, "Betriebsgrundsatz", "Datenbankänderungen werden vor dem Frontend ausgerollt und geprüft. Produktivtests vermeiden unbeabsichtigte Benachrichtigungen oder echte Planänderungen; vollständige Fachabläufe werden bei Bedarf in rückgerollten Transaktionen validiert.", BLUE_LIGHT)

    heading(doc, "10. Technischer Modul- und Datenbankanhang", 1)
    heading(doc, "10.1 Frontend-Module nach Domäne", 2)
    modules = [
        ("Authentifizierung und Sitzung", "supabase-auth, Passwort-Reset, Auth-Redirect, Workspace- und View-State"),
        ("Dienstplanung", "Schedule Core, Wochenboard, Monatsansicht, Vorlagen, SOLL-Summary, Zuweisung, Löschen, Verlauf"),
        ("Compliance", "Compliance Core, Workflow, UI, Plausibilitätsprüfung und PostgreSQL-Bridge"),
        ("Mitarbeiter", "Employee Core, Management V2, Team-Administration, Personalakte, Fristen und Erinnerungen"),
        ("Abwesenheiten", "Employee Workflow, Manager Workflow, Planning Guard und finale Verwaltungsansicht"),
        ("Arbeitszeit", "Time Tracking, UI Guard, Time Accounts, Feiertage, Monatsabschluss und Workspace"),
        ("Zusammenarbeit", "Benachrichtigungen, Schichttausch, Marktplatz und Störfall-Autopilot"),
        ("Auswertung", "Reports, Excel/PDF-Export und mehrstufiges PDF-Branding"),
        ("Oberfläche", "Landingpage, Sidebar, Topbar, Scrollbar, Hilfe-Center, Mobil-CSS und Mitarbeiter-Workspace"),
        ("Synchronisierung", "Supabase Data, Delta-Sync, Upsert-Guards, Legacy-Import und Delete-Bridge"),
    ]
    table(doc, ["Domäne", "Enthaltene Modulgruppen"], modules, [2500, 6860], 9)
    heading(doc, "10.2 Datenbankmigrationen", 2)
    db_rows = [
        ["shift_compliance.sql", "Compliance-Richtlinie, Zuweisungen, Änderungsanträge, Prüfungen, Freigaben und Audit."],
        ["company_user_management.sql", "Einladungen, Rollen- und Benutzerverwaltung."],
        ["security_hardening_v1.sql", "Least-Privilege-Grants und abgesicherte Funktionsausführung."],
        ["phase4_published_schedule_visibility.sql", "Sichtbarkeit veröffentlichter Dienstpläne."],
        ["legacy_assignment_review_v1.sql", "Prüfung und Abschluss migrierter Altzuweisungen."],
        ["shift_marketplace_v1.sql", "Angebot, Übernahme, Managerprüfung und Audit."],
        ["shift_marketplace_*_fix_v1.sql", "Portalrechte und Managerdatenfeed des Marktplatzes."],
        ["shift_readiness_v1.sql", "Schichtbestätigungen für die Einsatzbereitschaft."],
        ["disruption_autopilot_v1.sql", "Störfälle, Angebote, Kandidatenranking und atomare Übernahme."],
        ["grant_authenticated_time_entries_writes.sql", "Gezielte Rechte für Arbeitszeiteinträge."],
        ["fix_company_member_role_escalation.sql", "Blockiert unzulässige Rollenänderungen."],
    ]
    table(doc, ["Migration", "Zweck"], db_rows, [3300, 6060], 8.8)
    heading(doc, "10.3 Zentrale Datenobjekte", 2)
    for item in [
        "Unternehmen, Unternehmensmitglieder und Einladungen",
        "Mitarbeiter, Kontaktdaten, Freigaben und Personalinformationen",
        "Schichtvorlagen, SOLL-Besetzung, Dienstplan und Zuweisungen",
        "Abwesenheiten, Zeiteinträge, Stundenkonten und Monatsabschlüsse",
        "Änderungsanträge, Freigaben, Compliance-Läufe und Befunde",
        "Schichttausch, Marktplatzangebote, Störfälle und Ersatzangebote",
        "Benachrichtigungen, Bestätigungen und Audit-Ereignisse",
    ]: list_item(doc, item, bullet_id)
    heading(doc, "10.4 Begriffe", 2)
    table(doc, ["Begriff", "Bedeutung"], [
        ["SOLL", "Benötigte Anzahl oder Zielwert einer Schicht bzw. eines Zeitraums."],
        ["IST", "Tatsächlich zugewiesene Besetzung oder bestätigte Arbeitszeit."],
        ["RLS", "Row Level Security; datenbankseitige Sichtbarkeits- und Schreibregel je Datensatz."],
        ["RPC", "Serverseitig aufgerufene PostgreSQL-Funktion für einen kontrollierten Fachprozess."],
        ["Mandant", "Ein Unternehmen mit logisch getrennten Benutzern und Daten."],
        ["Audit", "Nachvollziehbares Ereignisprotokoll für fachlich relevante Änderungen."],
        ["Altbestand", "Migrierte ältere Schichtdaten, die gesondert geprüft und zugeordnet werden."],
        ["Einsatzbereitschaft", "Zusammengefasster Status aus Besetzung, Regeln, Veröffentlichung und Bestätigung."],
    ], [2200, 7160], 9.2)

    heading(doc, "11. Zusammenfassung und aktueller Stand", 1)
    para(doc, "SchichtFunk hat sich innerhalb der dokumentierten Aufbauphase von einem Dienstplanungsprototyp zu einer verknüpften Personal- und Einsatzplattform entwickelt. Der aktuelle Stand umfasst eine öffentliche Produktseite, ein vollständiges Verwaltungsportal, ein persönliches Mitarbeiterportal, cloudbasierte Datenhaltung, kontrollierte Automatisierung und abgesicherte Datenbankprozesse.")
    callout(doc, "Aktueller Produktkern", "Planen, prüfen, veröffentlichen, reagieren und auswerten - mit klar getrennten Rollen und einer gemeinsamen Datenbasis.")
    heading(doc, "11.1 Besonders differenzierende Funktionen", 2)
    for item in [
        "Einsatzbereitschafts-Ampel mit erklärbarem Status je Schicht.",
        "Störfall-Autopilot mit regelkonformem Kandidatenranking und erster sicherer Zusage.",
        "Kontrollierter Schicht-Marktplatz statt unprotokollierter Direktübernahme.",
        "Verknüpfung von Dienstplan, Abwesenheit, Arbeitszeit, Stundenkonto und Personalakte.",
        "Getrenntes Mitarbeiterportal im Unternehmenskontext mit persönlichen Einzelansichten.",
        "Datenbankseitige Mandanten- und Rollenabsicherung statt ausschließlichem UI-Schutz.",
    ]: list_item(doc, item, bullet_id)
    heading(doc, "11.2 Empfohlene Fortschreibung", 2)
    para(doc, "Diese Dokumentation sollte bei neuen Hauptfunktionen, Rollenänderungen, Datenbankmigrationen oder wesentlichen UI-Umbauten aktualisiert werden. Historische Meilensteine können fortlaufend ergänzt werden; technische Detailänderungen bleiben zusätzlich über Git-Historie und Migrationsdateien nachvollziehbar.")
    core = doc.core_properties
    core.title = "SchichtFunk Projekthistorie und Funktionsdokumentation"
    core.subject = "Hauptportal, Mitarbeiterportal, Funktionen, Architektur und Sicherheit"
    core.author = "SchichtFunk Projektteam"
    core.keywords = "SchichtFunk, Dienstplanung, Mitarbeiterportal, Supabase, Projektdokumentation"
    core.comments = "Erstellt aus dem dokumentierten Projektstand vom 31.08.2026."
    for current_table in doc.tables:
        if current_table.rows:
            set_repeat_table_header(current_table.rows[0])
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    build()
